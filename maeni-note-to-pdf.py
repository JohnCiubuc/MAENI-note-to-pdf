#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Created on Fri Aug  5 21:06:31 2022

@author: John Ciubuc
"""

import json
from docx import Document 
note = []

def noteSection(item, subItem = -1):
    if subItem == -1:   
        text = [x['content'] for x in note if x['item'] == item]
    else:
        text = [x['content'] for x in note if x['item'] == item and x['subitem'] == subItem]
    if len(text) == 0:
        return ''
    text = text[0].strip()
    return text

with open("../MAENI-dumps/student_note_content_202208052105.json", 'r') as f:
    jump = json.load(f)
    
note_db = jump['student_note_content']
note_ids = [329]
# note_ids = [329,359,389]

for id in note_ids:
    note =  [x for x in note_db if x['student_note_id'] ==id] 
    print(noteSection(20))
    # Add a heading of level 0 (Also called Title)
    document = Document()
    document.add_heading(f'Student Note ID: {id}', 0)
      
    document.add_heading('Chief Complaint', 1)
    document.add_paragraph(noteSection(20))
    document.add_heading('History of Presenting Illness', 1)
    document.add_paragraph(noteSection(21))
    document.add_heading('Review of Systems', 1)
    document.add_paragraph(noteSection(22))
    document.add_heading('History', 1)
    document.add_heading('Past Medical History', 2)
    document.add_paragraph(noteSection(23))
    document.add_heading('Past Surgical History', 2)
    document.add_paragraph(noteSection(24))
    document.add_heading('Medications', 2)
    document.add_paragraph(noteSection(25))
    document.add_heading('Allergies', 2)
    document.add_paragraph(noteSection(26))
    document.add_heading('Family History', 2)
    document.add_paragraph(noteSection(27))
    document.add_heading('Social History', 2)
    document.add_paragraph(noteSection(28))
    document.add_heading('Physical Exam', 1)
    document.add_heading('Vitals', 2)
    
    doc = f"""Heart Rate: {noteSection(29,0)}, Blood Pressure: {noteSection(29,1)}
Respiratory Rate: {noteSection(29,2)},  O2 Sat: {noteSection(29,3)}
Weight: {noteSection(29,4)}, Height: {noteSection(29,5)}"""
    
    document.add_paragraph(doc)
    document.add_heading('Exam', 2)
    document.add_paragraph(noteSection(39))
    document.add_heading('Data', 1)
    document.add_paragraph(noteSection(32))
    document.add_heading('Assessment and Plan', 1)
    document.add_heading('Summary Statement', 2)
    
    doc = f"""This is a {noteSection(38,1)} year old {noteSection(38,3)}, who is presenting today for {noteSection(38,6)}
The patient has a pertinent history of {noteSection(38,9)}
Patient's exam is remarkable for {noteSection(38,12)}
Patient's data is remarkable for {noteSection(38,15)}"""

    document.add_paragraph(doc)
    
    # Needs looping later
    problem_id=1
    document.add_heading(f'Problem {problem_id}', 3)
    document.add_heading('Differential DX}', 3)
    document.add_heading('Diagnostic Plan', 3)
    document.add_heading('Treatmnet Plan', 3)
    
    document.save('test.docx')